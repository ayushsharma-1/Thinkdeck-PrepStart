import os
import base64
import tempfile
from typing import Dict, Any
import PyPDF2
from docx import Document
from utils.logger import setup_logger
from utils.error_handler import handle_exceptions, ValidationError

logger = setup_logger(__name__)

class ResumeService:
    """Service for parsing resume files (PDF, DOCX)"""
    
    def __init__(self):
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        self.supported_types = ['pdf', 'docx', 'doc']
        
    @handle_exceptions
    async def parse_resume(
        self,
        file_data: str,
        file_name: str,
        file_type: str,
        session_id: str
    ) -> Dict[str, Any]:
        """Parse resume file and extract text content"""
        
        logger.info(f"Parsing resume for session: {session_id}, type: {file_type}")
        
        try:
            # Validate file type
            if file_type.lower() not in self.supported_types:
                raise ValidationError(f"Unsupported file type: {file_type}")
            
            # Decode base64 file data
            try:
                binary_data = base64.b64decode(file_data)
            except Exception as e:
                raise ValidationError(f"Invalid base64 file data: {str(e)}")
            
            # Validate file size
            if len(binary_data) > self.max_file_size:
                raise ValidationError(f"File too large: {len(binary_data)} bytes (max: {self.max_file_size})")
            
            if len(binary_data) == 0:
                raise ValidationError("File is empty")
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as temp_file:
                temp_file.write(binary_data)
                temp_file_path = temp_file.name
            
            try:
                # Parse based on file type
                if file_type.lower() == 'pdf':
                    text_content = await self._parse_pdf(temp_file_path)
                elif file_type.lower() in ['docx', 'doc']:
                    text_content = await self._parse_docx(temp_file_path)
                else:
                    raise ValidationError(f"Unsupported file type: {file_type}")
                
                # Validate extracted text
                if not text_content or len(text_content.strip()) == 0:
                    raise ValidationError("No text content found in the resume")
                
                # Clean and process text
                processed_text = self._clean_text(text_content)
                
                result = {
                    "success": True,
                    "text": processed_text,
                    "file_info": {
                        "name": file_name,
                        "type": file_type,
                        "size": len(binary_data),
                        "text_length": len(processed_text),
                        "session_id": session_id
                    }
                }
                
                logger.info(f"Resume parsed successfully for session: {session_id}, text length: {len(processed_text)}")
                return result
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                    
        except ValidationError:
            raise
        except Exception as e:
            logger.error(f"Error parsing resume for session {session_id}: {str(e)}")
            raise ValidationError(f"Resume parsing failed: {str(e)}")

    async def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file and extract text"""
        
        try:
            text_content = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise ValidationError("PDF file is encrypted and cannot be processed")
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n"
            
            if not text_content.strip():
                raise ValidationError("No text found in PDF file")
            
            return text_content
            
        except PyPDF2.errors.PdfReadError as e:
            raise ValidationError(f"Invalid PDF file: {str(e)}")
        except Exception as e:
            raise ValidationError(f"PDF parsing error: {str(e)}")

    async def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file and extract text"""
        
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            if not text_content.strip():
                raise ValidationError("No text found in DOCX file")
            
            return text_content
            
        except Exception as e:
            raise ValidationError(f"DOCX parsing error: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        
        try:
            # Remove excessive whitespace
            cleaned_text = "\n".join(line.strip() for line in text.split("\n") if line.strip())
            
            # Remove multiple consecutive newlines
            while "\n\n\n" in cleaned_text:
                cleaned_text = cleaned_text.replace("\n\n\n", "\n\n")
            
            # Remove special characters that might cause issues
            cleaned_text = cleaned_text.replace("\x00", "").replace("\r", "")
            
            return cleaned_text.strip()
            
        except Exception as e:
            logger.warning(f"Error cleaning text: {str(e)}")
            return text.strip()

    def validate_resume_content(self, text: str) -> Dict[str, Any]:
        """Validate and analyze resume content"""
        
        try:
            analysis = {
                "is_valid": False,
                "word_count": 0,
                "has_contact_info": False,
                "has_experience": False,
                "has_education": False,
                "has_skills": False,
                "warnings": []
            }
            
            if not text or len(text.strip()) < 50:
                analysis["warnings"].append("Resume content is too short")
                return analysis
            
            text_lower = text.lower()
            words = text.split()
            analysis["word_count"] = len(words)
            
            # Check for contact information
            contact_indicators = ['email', '@', 'phone', 'mobile', 'contact', 'linkedin']
            analysis["has_contact_info"] = any(indicator in text_lower for indicator in contact_indicators)
            
            # Check for experience section
            experience_indicators = ['experience', 'work', 'employment', 'career', 'position', 'job']
            analysis["has_experience"] = any(indicator in text_lower for indicator in experience_indicators)
            
            # Check for education
            education_indicators = ['education', 'university', 'college', 'degree', 'bachelor', 'master', 'phd']
            analysis["has_education"] = any(indicator in text_lower for indicator in education_indicators)
            
            # Check for skills
            skills_indicators = ['skills', 'technologies', 'programming', 'software', 'tools', 'languages']
            analysis["has_skills"] = any(indicator in text_lower for indicator in skills_indicators)
            
            # Determine if resume is valid
            valid_sections = sum([
                analysis["has_contact_info"],
                analysis["has_experience"],
                analysis["has_education"],
                analysis["has_skills"]
            ])
            
            analysis["is_valid"] = valid_sections >= 2 and analysis["word_count"] >= 100
            
            if not analysis["is_valid"]:
                analysis["warnings"].append("Resume may be missing important sections")
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error validating resume content: {str(e)}")
            return {"is_valid": False, "warnings": ["Content validation failed"]}

    def get_supported_formats(self) -> list:
        """Get list of supported resume formats"""
        
        return [
            {
                "format": "pdf",
                "mime_type": "application/pdf",
                "description": "PDF document",
                "max_size": f"{self.max_file_size // (1024*1024)}MB"
            },
            {
                "format": "docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "description": "Microsoft Word document",
                "max_size": f"{self.max_file_size // (1024*1024)}MB"
            }
        ]
